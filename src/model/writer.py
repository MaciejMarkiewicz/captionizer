from docx import Document
import datetime
import os


class Writer:

    def __init__(self, dest_dir):
        self._dest_dir = dest_dir
        self._document = Document()

    def write(self):
        self._document.add_heading('Transcription', 0)

        while True:
            text = (yield)
            self._document.add_paragraph(text)

    def save_current_document(self):
        date = datetime.datetime.now()
        filename = f'transcription{date.strftime("%Y-%m-%d_%H-%M-%S")}.docx'
        self._document.save(os.path.join(self._dest_dir, filename))
        self._document = Document()

        return filename
