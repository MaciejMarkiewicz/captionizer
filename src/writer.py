from docx import Document
import datetime
import os


class Writer:
    def write_result(self, result, dest_dir):
        document = Document()

        date = datetime.datetime.now()
        document.add_heading('Transcription', 0)
        document.add_paragraph(date.strftime("%d.%m.%Y, %H:%M:%S"))
        document.add_paragraph(result)

        filename = f'transcription{date.strftime("%Y-%m-%d_%H-%M-%S")}.docx'
        document.save(os.path.join(dest_dir, filename))

        return filename
